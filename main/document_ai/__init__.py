import logging
from flask_cors import CORS
import os
from flask import Flask, Blueprint, request, jsonify, send_from_directory, send_file
from jinja2 import Template
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Configure logging
logging.basicConfig(level=logging.ERROR)  # Set the log level as needed

# Create a blueprint
documentAI = Blueprint('documentAI', __name__)
CORS(documentAI)

# Define your routes
@documentAI.route('/')
def documentAI_route():
    return "I am Batman!"

# Sample rental agreement template
rental_agreement_template = """
**RENTAL AGREEMENT**

This Rental Agreement ("Agreement") is made and entered into this {{ date }} by and between:

**LANDLORD:**
Name: {{ landlord_name }}
Address: {{ landlord_address }}

**TENANT:**
Name: {{ tenant_name }}
Address: {{ tenant_address }}

**PROPERTY:**
The Landlord agrees to rent to the Tenant, and the Tenant agrees to rent from the Landlord, the following described property:

Address: {{ property_address }}
Description: {{ property_description }}

**TERMS AND CONDITIONS:**
1. **Rent:** The Tenant agrees to pay a monthly rent of {{ rent_amount }}, due on the {{ due_date }} of each month.

2. **Security Deposit:** A security deposit of {{ security_deposit_amount }} will be held by the Landlord to cover any damages or unpaid rent.

3. **Utilities:** The Tenant is responsible for paying all utilities and services, including {{ utilities }}.

4. **Maintenance:** The Tenant agrees to keep the property in good condition and promptly report any needed repairs to the Landlord.

5. **Term:** This Agreement shall commence on {{ start_date }} and continue on a month-to-month basis until terminated by either party with {{ notice_period }} days' written notice.

6. **Termination:** The Landlord or Tenant may terminate this Agreement with {{ notice_period }} days' written notice for any reason.

7. **Pets:** Pets are {{ pets_allowed }} on the property. Additional pet deposit/fees may apply.

8. **Insurance:** The Tenant is encouraged to obtain renter's insurance to cover personal belongings.

9. **Governing Law:** This Agreement shall be governed by the laws of {{ governing_law }}.

10. **Signatures:** This Agreement is binding upon both parties upon signing.

IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date first above written.

**LANDLORD:**
[Landlord Signature]                 [Landlord Name]           [Date]

**TENANT:**
[Tenant Signature]                   [Tenant Name]             [Date]
"""

# Define a route for generating a rental agreement
@documentAI.route('/api/generate_rental_agreement', methods=['POST'])
def generate_rental_agreement():
    try:
        # Get user input as JSON data from the POST request
        user_input = request.get_json()
        print('userInput', user_input)


        # Create a Jinja2 template from the rental agreement template
        template = Template(rental_agreement_template)

        # Render the template with user data
        rental_agreement_text = template.render(
            date=user_input.get('date', ''),
            landlord_name=user_input.get('landlordName', ''),
            landlord_address=user_input.get('landlordAddress', ''),
            tenant_name=user_input.get('tenantName', ''),
            tenant_address=user_input.get('tenantAddress', ''),
            property_address=user_input.get('propertyAddress', ''),
            property_description=user_input.get('propertyDescription', ''),
            rent_amount=user_input.get('rentAmount', ''),
            due_date=user_input.get('dueDate', ''),
            security_deposit_amount=user_input.get('securityDepositAmount', ''),
            utilities=user_input.get('utilities', ''),
            start_date=user_input.get('startDate', ''),
            notice_period=user_input.get('noticePeriod', ''),
            pets_allowed=user_input.get('petsAllowed', ''),
            governing_law=user_input.get('governingLaw', ''),
        )

        # Create a Word document
        doc = docx.Document()
        doc.add_heading('Rental Agreement', 0)

        # Add content to the Word document
        doc.add_paragraph(rental_agreement_text)

        # Save the Word document as a BytesIO object
        doc_bytesio = BytesIO()
        doc.save(doc_bytesio)
        doc_bytesio.seek(0)

        # Convert the Word document to a PDF
        pdf_bytesio = BytesIO()
        c = canvas.Canvas(pdf_bytesio, pagesize=letter)
        c.drawString(100, 750, 'Rental Agreement')
        text_object = c.beginText(100, 730)
        text_object.setFont("Helvetica-Bold", 14)
        text_object.textLines(rental_agreement_text)
        c.drawText(text_object)
        c.showPage()
        c.save()

        pdf_bytesio.seek(0)

        # Define the path to the "agreement" folder
        agreement_folder = 'agreement'

        # Ensure the "agreement" folder exists, or create it if it doesn't
        if not os.path.exists(agreement_folder):
            os.makedirs(agreement_folder)

        # Define the path for the saved PDF file in the "agreement" folder
        pdf_file_path = os.path.join(agreement_folder, 'rental_agreement.pdf')

        # Save the PDF file in the "agreement" folder
        with open(pdf_file_path, 'wb') as pdf_file:
            pdf_file.write(pdf_bytesio.read())

        # Return the path to the saved PDF file
        return jsonify({'downloadLink': pdf_file_path}), 200

    except Exception as e:
        logging.error(f'Error generating rental agreement: {str(e)}')
        return jsonify({'error': 'An error occurred while generating the rental agreement.'}), 500
